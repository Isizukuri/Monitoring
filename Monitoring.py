import time
from os import mkdir, chdir
from os.path import isdir, isfile

import requests
from BeautifulSoup import BeautifulSoup as BS
from urlparse import urljoin

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm

# user_name = raw_input('Введіть ваше прізвище та ініціали у форматі Прізвище І. П. через пробіли :')
# user_place = raw_input('Введіть вашу посаду :')
# search_word = raw_input('Введіть фрагмент тексту судового рішення: ')
#start_date = raw_input('Введіть початок періоду пошуку (дд.мм.рррр): ')
#end_date = raw_input('Введіть кінець періоду пошуку (дд.мм.рррр): ')

search_word = unicode('прокуратура', 'utf-8')
start_date = unicode('01.02.2015', 'utf-8')
end_date = unicode('28.02.2015', 'utf-8')
user_name = unicode('Маринич В. В.', 'utf-8')
user_place = unicode('прокурор відділу захисту прав і свобод дітей прокуратури Рівненської області', 'utf-8')


class Inputs(object):
    '''Description'''

    def __init__(self, start_date, end_date, search_word='',
                 court_name='166', court_region='18', cstype='2'):
        """

        :param Start date of the search perios:
        :param End date of the search period:
        :param Key word(s):
        """
        self.url = 'http://www.reyestr.court.gov.ua/'
        self.search_word = search_word
        self.start_date = start_date
        self.end_date = end_date
        self.search_parameters = {'SearchExpression': self.search_word,
                                  'CourtRegion[]': court_region,
                                  'CourtName[]': court_name,
                                  'RegDateBegin': self.start_date,
                                  'RegDateEnd': self.end_date,
                                  'CSType[]': cstype,
                                  'PagingInfo.ItemsPerPage': '25',
                                  'Liga': 'false'}
        self.requisites = {'texts': [], 'case_numbers': [], 'forms': [], 'dates': [], 'court_names': []}

    def __call__(self):
        """


        :raise RuntimeError: When the EDRSR is not avalaible
        """
        self.response = requests.post(self.url, self.search_parameters)
        if self.response.status_code != 200:
            raise RuntimeError('Got unexpected response in POST', self.response)
        else:
            print 'response success'

        self.cookies = self.response.cookies

        self.soup = BS(self.response.text)
        self.res_table = self.soup.find('table', id='tableresult')
        self.rows = self.res_table.findAll('tr')[1:]
        self.rel_links = [row.find('td').a['href'] for row in self.rows]

        #def savetofile(self):
        #if not isdir(self.path):
        #mkdir(self.path)
        #chdir(self.path)
        #self.path = ('./'+self.start_date+'-'+self.end_date)
        #Not Implemented

    def getAll(self):
        u"""Description"""
        self.child_page = ''
        self.text = ''
        for link in self.rel_links:
            time.sleep(5)
            self.child_page = BS(requests.get(urljoin(self.url, link)).text)
            self.requisites['texts'].append((self.child_page.body.find('textarea', id='txtdepository')).text)
        print 'Getall works'

        self.i = 1
        for row in self.rows:
            self.requisites['case_numbers'].append(row.find('td', {"class": "CaseNumber tr" + str(self.i)}).text)
            self.requisites['forms'].append(row.find('td', {"class": "CSType tr" + str(self.i)}).text)
            self.requisites['dates'].append(row.find('td', {"class": "RegDate tr" + str(self.i)}).text)
            self.requisites['court_names'].append(row.find('td', {"class": "CourtName tr" + str(self.i)}).text)
            if self.i == 1:
                self.i += 1
            else:
                self.i = 1
        return self.requisites


def outputs(requisites=None):
    """

    :param requisites: Take this from inputs.getAll()
    """
    dovidka = Document()

    global start_date
    global end_date
    global user_name
    uname = user_name.split(' ')

    section = dovidka.sections[-1]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    plain = dovidka.styles.add_style('Plain Text', WD_STYLE_TYPE.PARAGRAPH)
    plain.font.name = 'Times New Roman'
    plain.font.size = Pt(14)

    t_plain = dovidka.styles.add_style('Table Plain Text', WD_STYLE_TYPE.PARAGRAPH)
    t_plain.font.name = 'Times New Roman'
    t_plain.font.size = Pt(12)

    bold = dovidka.styles.add_style('Bold Text', WD_STYLE_TYPE.PARAGRAPH)
    bold.font.name = 'Times New Roman'
    bold.font.size = Pt(14)
    bold.font.bold = True

    t_bold = dovidka.styles.add_style('Table Bold Text', WD_STYLE_TYPE.PARAGRAPH)
    t_bold.font.name = 'Times New Roman'
    t_bold.font.size = Pt(12)
    t_bold.font.bold = True

    space = unicode(' ', 'utf-8')

    tmp = user_place.split(' ')
    x = 'прокурор'.decode('utf-8')
    tmp[tmp.index(x)] = unicode('прокурором', 'utf-8')
    try:
        if 'старший'.decode('utf-8') in tmp:
            x = 'старший'.decode('utf-8')
            tmp[tmp.index(x)] = unicode('старшим', 'utf-8')
    except ValueError: pass
    uplace = ''
    for i in tmp:
        uplace += (i+space)

    t1 = unicode("Мною, ", 'utf-8') + uplace
    t2 = unicode(
        " вивчено законність наступних судових рішень, занесених до Єдиного державного реєстру судових рішень, ",
        'utf-8')
    t3 = unicode("відібраних за датою з ", 'utf-8')
    t4 = unicode(' по ', 'utf-8')
    t5 = unicode(':', 'utf-8')
    next_line = unicode("\n", 'utf-8')

    headtext = unicode("ДОВІДКА \n про вивчення судових рішень \n Рівненського апеляційного господарського суду\n"
                       "та апеляційного суду Рівненської області \n за період з ",
                       'utf-8') + start_date + t4 + end_date

    heading1 = dovidka.add_paragraph(headtext, style='Bold Text')
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    text_before_table = t1 + uname[0] + space + uname[1] + uname[2] + t2 + t3 + start_date + t4 + end_date + t5

    paragraph1 = dovidka.add_paragraph(text_before_table, style='Plain Text')
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph1.paragraph_format.first_line_indent = Cm(1.25)
    paragraph1.paragraph_format.space_before = Pt(14)

    if requisites:
        table = dovidka.add_table(rows=1, cols=5, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_heading = table.rows[0]

        head_cell1 = table_heading.cells[0].add_paragraph(unicode('№ справи, дата рішення, суд', 'utf-8'),
                                                          style='Table Bold Text')
        head_cell2 = table_heading.cells[1].add_paragraph(unicode('Сторони', 'utf-8'),
                                                          style='Table Bold Text')
        head_cell3 = table_heading.cells[2].add_paragraph(unicode('Суть спору', 'utf-8'),
                                                          style='Table Bold Text')
        head_cell4 = table_heading.cells[3].add_paragraph(unicode(
            "Ціна позову, площа земель, інші дані, які характеризують правовідносини", 'utf-8'),
            style="Table Bold Text")
        head_cell5 = table_heading.cells[4].add_paragraph(unicode("Вжиті заходи або висновок про законність", 'utf-8'),
                                                          style="Table Bold Text")

        for i in range(len(requisites['forms'])):
            table.add_row()

        work_rows = table.rows[1:]

        for row in work_rows:
            cell = row.cells[0]
            cell_contents = cell.add_paragraph((requisites['case_numbers'][work_rows.index(row)] + ', ',
                                                requisites['dates'][work_rows.index(row)] + ', ',
                                                requisites['court_names'][work_rows.index(row)]),
                                               style='Table Plain Text')

    text_after_table = unicode('За результатами моніторингу Єдиного державного реєстру '
                                           'судових рішень за вказаний період незаконних судових рішень у цивільних'
                                           ' справах з питань захисту прав дітей (про позбавлення батьківських прав, '
                                           'відібрання дитини без позбавлення батьківських прав, усиновлення дітей іноземними '
                                           'громадянами тощо), а також у господарських справах з питань захисту інтересів держави '
                                           'у сфері охорони дитинства, винесених без участі прокурорів, не виявлено.',
                                           'utf-8')

    paragraph2 = dovidka.add_paragraph(text_after_table, style='Plain Text')
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph2.paragraph_format.first_line_indent = Cm(1.25)
    paragraph1.paragraph_format.space_after = Pt(14)

    signment = dovidka.add_table(rows=1, cols=2)
    signment.columns[0].width = Cm(12.5)
    signment.columns[1].width = Cm(5.4)
    sign_prefix = signment.rows[0].cells[0].add_paragraph(user_place.capitalize(), style='Bold Text')
    sign_prefix.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sign_end = signment.rows[0].cells[1].add_paragraph(next_line + next_line + uname[1] + space + uname[0],
                                                       style='Bold Text')
    sign_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    dovidka.save('dovidka.docx')
    
        
        

        


